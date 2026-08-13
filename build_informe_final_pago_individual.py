from pathlib import Path
import sqlite3, json
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\Ingenieria\VigIA")
DB = ROOT / "datos" / "premio_productividad.db"
OUT = ROOT / "outputs" / "informe_final_pago_individual_real"
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / "analisis_final_pago_individual_bultos_reales.docx"
RUN = "P0_4f8ad82ed081477f_20260807175215"
NAVY, BLUE, TEAL, INK, MUTED = "123047", "2E74B5", "0F766E", "1F2937", "5B6573"
LIGHT, GREEN, GOLD = "F2F4F7", "1F6B4F", "8A6400"

def money(v): return f"${v:,.0f}".replace(",", ".")
def pct(v): return f"{v:.1f}%".replace(".", ",")
def run_style(run, size=10.5, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"; run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri"); run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size); run.font.color.rgb = RGBColor.from_string(color); run.bold = bold; run.italic = italic

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn("w:shd"))
    if shd is None: shd = OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def cell_margins(cell):
    tcPr = cell._tc.get_or_add_tcPr(); mar = tcPr.find(qn("w:tcMar"))
    if mar is None: mar = OxmlElement("w:tcMar"); tcPr.append(mar)
    for side, val in [("top",80),("start",120),("bottom",80),("end",120)]:
        el = mar.find(qn(f"w:{side}"))
        if el is None: el = OxmlElement(f"w:{side}"); mar.append(el)
        el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa")

def geometry(table, widths, indent=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT; table.autofit = False
    pr = table._tbl.tblPr
    tw = pr.find(qn("w:tblW"))
    if tw is None: tw = OxmlElement("w:tblW"); pr.append(tw)
    tw.set(qn("w:w"), str(sum(widths))); tw.set(qn("w:type"), "dxa")
    ti = pr.find(qn("w:tblInd"))
    if ti is None: ti = OxmlElement("w:tblInd"); pr.append(ti)
    ti.set(qn("w:w"), str(indent)); ti.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for item in list(grid): grid.remove(item)
    for w in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(w)); grid.append(col)
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            pr = cell._tc.get_or_add_tcPr(); tcw = pr.find(qn("w:tcW"))
            if tcw is None: tcw = OxmlElement("w:tcW"); pr.append(tcw)
            tcw.set(qn("w:w"), str(w)); tcw.set(qn("w:type"), "dxa"); cell_margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def repeat_header(row):
    trPr = row._tr.get_or_add_trPr(); el = OxmlElement("w:tblHeader"); el.set(qn("w:val"), "true"); trPr.append(el)

def para(doc, text, size=10.5, color=INK, bold=False, italic=False, align=None, after=6):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = 1.1
    if align is not None: p.alignment = align
    run_style(p.add_run(text), size, color, bold, italic); return p

def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}"); p.paragraph_format.keep_with_next = True
    run_style(p.add_run(text), {1:16,2:13,3:12}[level], BLUE if level < 3 else NAVY, True); return p

def bullets(doc, items):
    for text in items:
        p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.left_indent = Inches(.5); p.paragraph_format.first_line_indent = Inches(-.25); p.paragraph_format.space_after = Pt(4)
        run_style(p.add_run(text), 10.5)

def table(doc, headers, rows, widths, money_cols=(), pct_cols=(), font=8.8):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; geometry(t, widths); repeat_header(t.rows[0])
    for c, h in zip(t.rows[0].cells, headers):
        shade(c, NAVY); p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0); run_style(p.add_run(str(h)), font, "FFFFFF", True)
    for ri, row in enumerate(rows):
        cells=t.add_row().cells
        for i, v in enumerate(row):
            if i in money_cols: txt=money(float(v))
            elif i in pct_cols: txt=pct(float(v))
            elif isinstance(v, float): txt=f"{v:,.1f}".replace(",", ".")
            else: txt=str(v)
            p=cells[i].paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT if i in money_cols or i in pct_cols else WD_ALIGN_PARAGRAPH.LEFT; run_style(p.add_run(txt), font)
            if ri % 2: shade(cells[i], LIGHT)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

def callout(doc, title, text, fill="E8EEF5", accent=TEAL):
    t=doc.add_table(rows=1, cols=1); geometry(t,[9360]); c=t.cell(0,0); shade(c,fill)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(3); run_style(p.add_run(title),10.5,accent,True)
    p=c.add_paragraph(); p.paragraph_format.space_after=Pt(0); run_style(p.add_run(text),10.3)
    doc.add_paragraph().paragraph_format.space_after=Pt(3)

def page_field(p):
    r=p.add_run(); a=OxmlElement("w:fldChar"); a.set(qn("w:fldCharType"),"begin"); b=OxmlElement("w:instrText"); b.set(qn("xml:space"),"preserve"); b.text=" PAGE "; c=OxmlElement("w:fldChar"); c.set(qn("w:fldCharType"),"end"); r._r.append(a); r._r.append(b); r._r.append(c); run_style(r,8.5,MUTED)

con=sqlite3.connect(DB)
p0=pd.read_sql_query("select * from pp_punto0_legajo_dia where run_id=?",con,params=(RUN,))
pay=pd.read_sql_query("select * from pp_evaluacion_picking_pago where query_version='evaluacion_picking_pago_real_picking_v10_calidad_tnc_error' and fecha_base between '2026-07-01' and '2026-07-31'",con)
pay["factor_multiplicador"]=pay["factor_multiplicador"].fillna(1).astype(float)
p0=p0.merge(pay[["fecha_base","legajo","factor_multiplicador"]],on=["fecha_base","legajo"],how="left"); p0["factor_multiplicador"]=p0["factor_multiplicador"].fillna(1); p0["nuevo_ajustado"]=p0["premio_individual"]*p0["factor_multiplicador"]
a=p0.groupby("legajo").agg(actual=("premio_actual","sum"),nuevo_base=("premio_individual","sum"),nuevo=("nuevo_ajustado","sum"),bultos=("bultos_reales","sum"),horas=("horas","sum"),dias=("fecha_base","nunique"),sector=("sector",lambda s:s.value_counts().index[0])).reset_index(); a["bph"]=a["bultos"]/a["horas"]; a["dif_pct"]=(a["nuevo"]/a["actual"]-1)*100
g=pay.groupby("legajo").agg(bultos_actuales=("bultos_actuales","sum"),equivalentes=("total_equivalentes","sum"),equiv_sector=("equivalencia_sector","sum"),equiv_traslado=("equivalencia_traslado","sum"),equiv_consol=("equivalencia_consolidacion","sum"),nivel_actual=("nivel_actual","mean"),dias_factor2=("factor_multiplicador",lambda s:(s==2).sum()),tnc=("penalizacion_tnc","sum"),error=("penalizacion_error","sum")).reset_index(); g["equiv_adicional"]=g["equivalentes"]-g["bultos_actuales"]; g["pct_equiv"]=np.where(g["equivalentes"]>0,g["equiv_adicional"]/g["equivalentes"]*100,0); a=a.merge(g,on="legajo",how="left")
meta=json.load(open(ROOT/"datos"/"premio_tabla_foto.json",encoding="utf-8")); sm={str(x["sector"]):x for x in meta["sectores"]}; a["grupo"]=a["sector"].map(lambda s:sm.get(str(s),{}).get("grupo_productivo",""))
positive=a[a["actual"]>0].copy(); total_actual=float(a["actual"].sum()); total_new=float(a["nuevo"].sum()); diff=total_new-total_actual
sec=positive.groupby(["sector","grupo"]).agg(legajos=("legajo","count"),perdedores=("dif_pct",lambda x:(x<0).sum()),actual=("actual","sum"),nuevo=("nuevo","sum"),eq_adicional=("equiv_adicional","sum"),pct_equiv=("pct_equiv","mean"),bph=("bph","mean")).reset_index(); sec["dif_pct"]=(sec["nuevo"]/sec["actual"]-1)*100; sec["perdida"]=sec["actual"]-sec["nuevo"]; sec=sec.sort_values("perdida",ascending=False)
segments=[]
for lo,hi,label in [(-1e9,-60,"60% o más menos"),(-60,-30,"30% a 60% menos"),(-30,-10,"10% a 30% menos"),(-10,0,"Hasta 10% menos"),(0,10,"Hasta 10% más"),(10,30,"10% a 30% más"),(30,60,"30% a 60% más"),(60,1e9,"Más de 60% más")]:
    x=positive[(positive["dif_pct"]>lo)&(positive["dif_pct"]<=hi)]; segments.append((label,len(x),len(x)/len(positive)*100,x["actual"].sum(),x["nuevo"].sum()))
examples=[]
for leg in ["203637","207397","206693","206101","735249"]:
    r=a[a.legajo.astype(str)==leg].iloc[0]; examples.append((leg,r.sector,r.actual,r.nuevo_base,r.nuevo,r.dif_pct,r.bph,r.dias,r.horas,r.equiv_adicional,r.pct_equiv,r.nivel_actual,r.dias_factor2))
con.close()

doc=Document(); s=doc.sections[0]; s.page_width=Inches(8.5); s.page_height=Inches(11); s.top_margin=s.bottom_margin=s.left_margin=s.right_margin=Inches(1); s.header_distance=s.footer_distance=Inches(.492)
for name,size,color,before,after in [("Normal",11,INK,0,6),("Heading 1",16,BLUE,16,8),("Heading 2",13,BLUE,12,6),("Heading 3",12,NAVY,8,4)]:
    st=doc.styles[name]; st.font.name="Calibri"; st._element.rPr.rFonts.set(qn("w:ascii"),"Calibri"); st._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri"); st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color); st.font.bold=name!="Normal"; st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.line_spacing=1.1
hdr=s.header.paragraphs[0]; run_style(hdr.add_run("ANÁLISIS DE PAGO INDIVIDUAL  |  PICKING"),8.5,MUTED,True)
f=s.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.RIGHT; run_style(f.add_run("VigIA  ·  Julio 2026  ·  Página "),8.5,MUTED); page_field(f)
para(doc,"INFORME DE DECISIÓN",9.5,TEAL,True,after=8); p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4); run_style(p.add_run("Pago individual por bultos reales"),25,NAVY,True); para(doc,"Análisis de impacto del nuevo método y propuestas de ajuste sin equivalencias",13,MUTED,after=14); para(doc,"Período: 1 al 31 de julio de 2026  |  Universo: Picking  |  Snapshot de cálculo congelado",9.5,MUTED,after=12)
callout(doc,"Premisa de dirección","El pago nuevo debe explicarse con bultos reales. Metros, traslados y equivalencias se utilizan solo para diagnosticar dónde la escala necesita calibración; no se incorporan como componente de pago.","EAF4F1",TEAL)

heading(doc,"1. Resumen ejecutivo")
para(doc,f"El pago actual suma {money(total_actual)} y el pago individual nuevo, aplicando el multiplicador del día, suma {money(total_new)}. La diferencia es {money(diff)} ({pct(diff/total_actual*100)}). De los 252 legajos con pago actual positivo, 159 cobran menos y 93 cobran más.")
table(doc,["Indicador","Resultado"],[["Pago actual total",money(total_actual)],["Pago individual nuevo ajustado",money(total_new)],["Diferencia total",f"{money(diff)} ({pct(diff/total_actual*100)})"],["Legajos con merma","159 de 252 (63,1%)"],["Merma de 30% o más","57 legajos (22,6%)"],["TNC + errores",money(pay["penalizacion_tnc"].sum()+pay["penalizacion_error"].sum())]],[3900,5460])
callout(doc,"Mensaje central","El multiplicador corrige una parte de la diferencia, pero queda una merma concentrada en sectores donde el modelo actual reconoce distancia y complejidad mediante bultos equivalentes y la escala nueva observa únicamente bultos reales por hora.","FFF8E8",GOLD)

heading(doc,"2. Alcance y método")
para(doc,"La comparación es individual contra individual. No incluye premio grupal ni adicionales grupales. El nuevo pago se calcula sobre bultos reales por hora y se ajusta con el multiplicador vigente para cada legajo y día.")
bullets(doc,["Base actual: pago real Picking con nivel, bultos actuales, total de equivalentes, multiplicador, TNC y errores.","Base nueva: premio individual de la escala sectorial horaria sobre bultos reales.","Diagnóstico: diferencia entre bultos actuales y total de equivalentes, separando sector, traslado y consolidación.","Regla: la equivalencia explica el origen operativo de la brecha, pero no se propone pagarla."])

heading(doc,"3. Distribución del impacto")
table(doc,["Segmento","Legajos","% legajos","Pago actual","Pago nuevo"],segments,[2700,900,1100,2300,2360],money_cols={3,4},pct_cols={2})
para(doc,"La mayor concentración está en mermas de 10% a 30%, pero 57 legajos presentan una reducción de al menos 30%. El impacto depende de sector, nivel actual, multiplicador y peso de las equivalencias.",9.5,MUTED,italic=True)

heading(doc,"4. Qué explica la brecha")
heading(doc,"4.1. Multiplicador",2)
para(doc,"Hay 121 jornadas con multiplicador 2. Sin trasladarlo al nuevo pago horario, esas jornadas quedaban aproximadamente 70,2% por debajo del actual. Al incorporarlo, el subconjunto queda 19,1% por encima en promedio. Debe formar parte del nuevo cálculo, con una regla explícita y auditable.")
heading(doc,"4.2. Equivalencias como diagnóstico de distancia y complejidad",2)
pay_equiv_adicional=pay["total_equivalentes"]-pay["bultos_actuales"]
para(doc,f"En el cache de pago hay {(pay_equiv_adicional>0).sum():,} registros con equivalencias adicionales positivas. Los bultos actuales suman {pay['bultos_actuales'].sum():,.0f} y el total equivalente {pay['total_equivalentes'].sum():,.0f}; la diferencia es {pay_equiv_adicional.sum():,.0f} unidades equivalentes. En AM, PI, VA y N1 las equivalencias representan aproximadamente entre 34% y 38% del total equivalente.")
table(doc,["Sector / grupo","Legajos","Perdedores","Equiv. adicionales","% equiv.","Merma ajustada"],[(f"{r.sector} / {r.grupo}",int(r.legajos),int(r.perdedores),r.eq_adicional,r.pct_equiv,r.dif_pct) for _,r in sec.head(10).iterrows()],[1900,850,1000,1900,1400,2310],money_cols={3},pct_cols={4,5})
para(doc,"La evidencia permite decir que ciertos sectores tenían una carga operativa reconocida por equivalencias. La solución propuesta no es volver a pagar equivalencias: es ajustar la escala de bultos reales de esos sectores.")

heading(doc,"4.3. Sectores con mayor impacto")
table(doc,["Sector / grupo","Legajos","Perdedores","Actual","Nuevo","Diferencia"],[(f"{r.sector} / {r.grupo}",int(r.legajos),int(r.perdedores),r.actual,r.nuevo,r.dif_pct) for _,r in sec.head(10).iterrows()],[1900,750,900,1900,1900,2010],money_cols={3,4},pct_cols={5})
para(doc,"B1 concentra la mayor pérdida monetaria por volumen. AM concentra la mayor merma relativa. PI, N1 y VA también presentan una brecha elevada con pocos legajos, por lo que requieren revisión específica de la tabla.",9.5,MUTED,italic=True)

heading(doc,"5. Casos reales")
table(doc,["Legajo","Sector","Actual","Nuevo base","Nuevo ajust.","Dif.","Bultos/h","Equiv. adic.","Nivel actual"],[(e[0],e[1],e[2],e[3],e[4],e[5],e[6],e[9],e[11]) for e in examples],[800,700,1300,1350,1350,850,950,1400,660],money_cols={2,3,4},pct_cols={5})
para(doc,"203637 mejora al aplicar el multiplicador, pero permanece 45,1% debajo del pago actual. 207397 y 206693 mantienen mermas superiores al 68% aun con el ajuste. Estos casos muestran que el multiplicador no resuelve la diferencia de escala.")

heading(doc,"6. Propuestas sin equivalencias")
callout(doc,"Propuesta recomendada: escala sectorial de bultos reales","Mantener una unidad única y explicable —el bulto real—, pero definir por sector cortes de bultos/hora alcanzables y premios por nivel calibrados con la productividad real observada. El sector no agrega bultos: modifica el umbral y/o el valor porque la operación es distinta.","EAF4F1",GREEN)
table(doc,["Alternativa","Qué cambia","Ventaja","Control necesario"],[
["A. Bajar umbrales","Los niveles superiores requieren menos bultos reales/hora en sectores lentos por layout.","Compensa el acceso a premios sin inventar producción.","Validar con percentiles y costo máximo."],
["B. Subir premios","Se mantienen cortes de bultos reales, pero cambia el valor del premio en sectores seleccionados.","Muy simple: mismo bulto real, tarifa sectorial distinta.","Evitar sobrepago y saltos injustificados."],
["C. Escala híbrida","Baja moderada de umbrales y ajuste moderado del premio.","Reparte la corrección entre acceso y valor.","Simular por legajo y sector."],
["D. Transición","Piso temporal de merma mientras se valida la tabla nueva.","Reduce conflicto de implementación.","Debe tener fecha de vencimiento."]],[1800,2450,2400,2710],font=8.4)
heading(doc,"6.1. Criterio de recomendación",2)
bullets(doc,["Ajustar primero AM, PI, N1 y VA, donde la merma relativa es alta y la proporción de equivalencias diagnósticas también.","Revisar B1 por impacto total: la merma porcentual es menor, pero concentra más personas y pesos.","Mantener el multiplicador dentro del pago horario nuevo.","Usar las equivalencias para identificar sectores, nunca para calcular el pago.","Simular cada alternativa con monto total, legajos beneficiados, mermas restantes y posibles sobrepagos."])

heading(doc,"7. Proceso de cierre")
table(doc,["Paso","Resultado esperado"],[["1. Diagnóstico","Mapa de sectores donde las equivalencias explican saltos de escala actual."],["2. Calibración","Nueva tabla sectorial basada exclusivamente en bultos reales/hora."],["3. Simulación","Comparación actual vs. nuevo ajustado, sin grupal, por legajo y sector."],["4. Control","Mermas mayores a 10%, 30% y 50%; sobrepagos; costo total."],["5. Decisión","Elegir umbrales, premios y transición con impacto explícito."]],[1400,7960])

heading(doc,"8. Conclusión")
para(doc,"El nuevo sistema puede cumplir la premisa de pagar por bultos reales sin ignorar la realidad operativa de los sectores. La distancia y el recorrido no deben convertirse en bultos ficticios ni en equivalencias pagables; deben servir para detectar que algunos sectores necesitan una escala propia de bultos reales. La recomendación es calibrar umbrales y premios por sector, conservar el multiplicador de jornada y validar el resultado por legajo antes de reemplazar el esquema vigente.")
heading(doc,"Anexo — definiciones y límites")
table(doc,["Concepto","Definición"],[["Bultos reales","Unidades físicas registradas para el legajo."],["Bultos equivalentes","Variable del pago actual que combina bultos y componentes de sector, traslado o consolidación; solo diagnóstico."],["Pago nuevo ajustado","Premio individual horario nuevo multiplicado por el factor vigente del día."],["Merma","Pago nuevo ajustado menos pago actual, como porcentaje del pago actual."],["Límite","Las equivalencias detectan el origen operativo de la brecha, pero no forman parte de la fórmula propuesta."]],[2100,7260])
para(doc,"Fuente: cache local del módulo de análisis de productividad, snapshot de tablas de premios Picking y tablas de pago/evaluación de julio 2026. Documento preparado como base de decisión; los valores de las alternativas deben cerrarse con una simulación parametrizada antes de producción.",8.5,MUTED,italic=True,after=0)
doc.core_properties.title="Análisis final de pago individual por bultos reales"; doc.core_properties.subject="Impacto y propuestas sin equivalencias"; doc.core_properties.author="VigIA"; doc.save(DOCX)
print(DOCX)
