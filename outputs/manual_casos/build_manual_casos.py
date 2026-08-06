from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
ROOT = BASE.parents[1]
SCREENSHOTS = BASE / "screenshots"
OUT = BASE / "Manual_Usuario_VigIA_Casos_Racks.docx"
APP_LOGO = ROOT / "resources" / "APPLogo.png"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(90, 103, 114)
GRID = "CBD5E1"
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
GREEN = RGBColor(24, 128, 82)
RED = RGBColor(160, 32, 32)
AMBER = RGBColor(145, 91, 0)

SPANISH_REPLACEMENTS = {
    "Indice": "Índice",
    "Modulo": "Módulo",
    "Reparacion": "Reparación",
    "Gestion": "Gestión",
    "Ambito": "Ámbito",
    "Version": "Versión",
    "Pagina": "Página",
    "Operacion": "Operación",
    "operacion": "operación",
    "situacion": "situación",
    "informacion": "información",
    "minima": "mínima",
    "accion": "acción",
    "validacion": "validación",
    "Validacion": "Validación",
    "posicion": "posición",
    "Posicion": "Posición",
    "ubicacion": "ubicación",
    "inutilizacion": "inutilización",
    "Inutilizacion": "Inutilización",
    "rehabilitacion": "rehabilitación",
    "fisicamente": "físicamente",
    "correccion": "corrección",
    "descripcion": "descripción",
    "critica": "crítica",
    "tecnico": "técnico",
    "estan": "están",
    "areas": "áreas",
    "numero": "número",
    "acompanhar": "acompañar",
    "Acompanhar": "Acompañar",
}


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=GRID, sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_in):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            row.cells[idx].width = Inches(width)
            set_cell_border(row.cells[idx])
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.TOP


def set_paragraph_border_bottom(paragraph, color="CBD5E1", size="8"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def add_page_number(paragraph):
    run = paragraph.add_run("Página ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(fld_begin)
    run = paragraph.add_run()
    run._r.append(instr)
    run = paragraph.add_run()
    run._r.append(fld_sep)
    run = paragraph.add_run()
    run._r.append(fld_end)


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.25


def setup_header_footer(section):
    header = section.header.paragraphs[0]
    header.text = ""
    r = header.add_run("VigIA - Manual de Usuario | Casos de Reparacion de Racks")
    set_run_font(r, size=9, color=MUTED, bold=True)
    set_paragraph_border_bottom(header, color="D5DEE8", size="6")

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left = footer.add_run("VigIA v2.0 | CD Coto | Uso interno")
    set_run_font(left, size=9, color=MUTED)
    sep = footer.add_run(" | ")
    set_run_font(sep, size=9, color=MUTED)
    add_page_number(footer)


def add_title_block(doc):
    if APP_LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(APP_LOGO), width=Inches(1.05))
        p.paragraph_format.space_after = Pt(28)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("VigIA")
    set_run_font(r, size=30, color=INK, bold=True)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Manual de Usuario")
    set_run_font(r, size=22, color=BLUE, bold=True)
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Modulo Casos - Reparacion de Racks")
    set_run_font(r, size=15, color=MUTED)
    p.paragraph_format.space_after = Pt(26)

    add_metadata_table(
        doc,
        [
            ("Destinatarios", "ADO, Mapa, Mantenimiento e Ingenieria"),
            ("Ambito", "Gestion operativa de services de racks, ubicaciones y adjuntos"),
            ("Version", "1.0"),
            ("Fecha", date.today().strftime("%Y-%m-%d")),
        ],
        widths=(1.875, 4.625),
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Documento de capacitacion y consulta para usuarios finales.")
    set_run_font(r, size=11, color=MUTED, italic=True)
    p.paragraph_format.space_before = Pt(18)
    doc.add_page_break()


def add_metadata_table(doc, rows, widths=(1.875, 4.625)):
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table, widths)
    for i, (label, value) in enumerate(rows):
        set_cell_shading(table.cell(i, 0), HEADER_FILL)
        p = table.cell(i, 0).paragraphs[0]
        p.text = ""
        r = p.add_run(label)
        set_run_font(r, size=9, color=INK, bold=True)
        p = table.cell(i, 1).paragraphs[0]
        p.text = ""
        r = p.add_run(value)
        set_run_font(r, size=10.5, color=INK)
    doc.add_paragraph()


def add_callout(doc, title, body, color=INK):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, (6.5,))
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.text = ""
    r = p.add_run(title)
    set_run_font(r, size=10, color=color, bold=True)
    p.add_run("\n")
    r = p.add_run(body)
    set_run_font(r, size=10, color=INK)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("• ")
        set_run_font(r, size=11, color=BLUE, bold=True)
        r = p.add_run(item)
        set_run_font(r, size=11, color=INK)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=9, color=MUTED, italic=True)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)


def add_image(doc, name, caption, width=6.45):
    path = SCREENSHOTS / name
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        add_caption(doc, caption)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.text = ""
        r = p.add_run(header)
        set_run_font(r, size=9, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.text = ""
            r = p.add_run(value)
            set_run_font(r, size=9.5, color=INK)
            set_cell_border(cells[idx])
            set_cell_margins(cells[idx])
    doc.add_paragraph()


def polish_spanish(doc):
    replacements = sorted(SPANISH_REPLACEMENTS.items(), key=lambda item: len(item[0]), reverse=True)

    def replace_words(text):
        for before, after in replacements:
            text = re.sub(rf"\b{re.escape(before)}\b", after, text)
        return text

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text:
                run.text = replace_words(run.text)
    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                for run in paragraph.runs:
                    if run.text:
                        run.text = replace_words(run.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text:
                            run.text = replace_words(run.text)


def add_toc(doc):
    doc.add_heading("Índice de Contenido", level=1)
    sections = [
        "1. Objetivo y alcance",
        "2. Roles y responsabilidades",
        "3. Flujo operativo",
        "4. Pantallas principales",
        "5. Trabajo por rol",
        "6. Estados, acciones y reglas",
        "7. Ingresos desde Forms",
        "8. Buenas practicas y soporte",
    ]
    add_bullets(doc, sections)
    add_callout(
        doc,
        "Nota",
        "El manual está pensado para capacitación operativa. Las capturas corresponden a la app abierta en Chrome al momento de generar este documento.",
    )
    doc.add_page_break()


def build():
    doc = Document()
    style_doc(doc)
    setup_header_footer(doc.sections[0])
    add_title_block(doc)
    add_toc(doc)

    doc.add_heading("1. Objetivo y alcance", level=1)
    doc.add_paragraph(
        "Este manual explica el uso del modulo de casos de VigIA para gestionar reparaciones de racks, "
        "desde el ingreso por Microsoft Forms hasta el cierre operativo con ubicaciones rehabilitadas en WMS."
    )
    add_bullets(
        doc,
        [
            "Unificar la lectura del caso entre ADO, Mapa, Mantenimiento e Ingenieria.",
            "Asegurar que cada etapa avance con la informacion minima requerida.",
            "Facilitar el seguimiento de adjuntos, ubicaciones, service externo y observaciones.",
            "Detectar diferencias entre servicios activos y ubicaciones inutilizadas en WMS.",
        ],
    )

    doc.add_heading("2. Roles y responsabilidades", level=1)
    add_matrix(
        doc,
        ["Rol", "Responsabilidad principal", "Pantallas clave"],
        [
            ("Operacion", "Carga la solicitud desde Forms y aporta informacion inicial del rack.", "Forms, Bandeja, Detalle"),
            ("ADO", "Valida posiciones y registra el service externo antes de derivar a Mapa.", "Bandeja Kanban, Detalle"),
            ("Mapa", "Gestiona traspasos WMS, inutiliza/rehabilita ubicaciones y cierra el caso.", "Bandeja Kanban, Control de Ubicaciones"),
            ("Mantenimiento", "Ejecuta la reparacion del rack asociada al service externo.", "Bandeja Kanban, Detalle"),
            ("Ingenieria", "Da soporte al flujo, monitorea indicadores, ingresos Forms y diferencias WMS.", "Dashboard, Control de Ubicaciones, Ingresos Forms"),
        ],
        (1.25, 3.25, 2.0),
    )

    doc.add_heading("3. Flujo operativo", level=1)
    doc.add_paragraph(
        "La solapa Flujo muestra el circuito completo del caso. El objetivo es capacitar al usuario sobre "
        "quien interviene, que valida y cual es la siguiente accion esperada."
    )
    add_image(doc, "02_flujo.png", "Solapa Flujo: secuencia operativa del caso.")
    add_matrix(
        doc,
        ["Paso", "Responsable", "Criterio de avance"],
        [
            ("1", "Operacion", "Ingresa el caso mediante Forms."),
            ("2", "ADO", "Valida que las posiciones existan. Si no existen, rechaza o solicita correccion."),
            ("3", "ADO", "Carga el numero de service externo; sin este dato no se deriva a Mapa."),
            ("4", "Mapa", "Genera traspasos WMS y deja observaciones operativas."),
            ("5", "Mapa", "Confirma traspasos finalizados y verifica si las ubicaciones quedaron vacias."),
            ("6", "Mapa", "Inutiliza las ubicaciones en WMS cuando corresponde."),
            ("7", "Mantenimiento", "Realiza la reparacion asociada al service externo."),
            ("8", "Mantenimiento", "Finaliza su intervencion y deriva a Mapa."),
            ("9", "Mapa", "Releva fisicamente la ubicacion y reetiqueta si hace falta."),
            ("10", "Mapa", "Quita la inutilizacion en WMS y cierra el caso."),
        ],
        (0.65, 1.35, 4.5),
    )

    doc.add_heading("4. Pantallas principales", level=1)
    doc.add_heading("Dashboard", level=2)
    doc.add_paragraph(
        "Resume la situacion operativa: tickets abiertos por sector, servicios por zona, tickets por estado "
        "y control WMS versus services."
    )
    add_image(doc, "01_dashboard.png", "Dashboard operativo con indicadores y graficos.")

    doc.add_heading("Control de Ubicaciones", level=2)
    doc.add_paragraph(
        "Permite cruzar ubicaciones inutilizadas en WMS contra services activos cargados en VigIA. "
        "Es la pantalla de control para detectar ubicaciones inutilizadas sin service y services sin bloqueo WMS."
    )
    add_image(doc, "03_control_ubicaciones.png", "Control de Ubicaciones: cruce WMS versus VigIA.")

    doc.add_heading("Bandeja Kanban", level=2)
    doc.add_paragraph(
        "Es la vista recomendada para trabajar. Cada columna representa una etapa o responsable del flujo, "
        "y cada tarjeta corresponde a un caso."
    )
    add_image(doc, "04_bandeja_kanban.png", "Bandeja Kanban: seguimiento por columnas operativas.")
    add_bullets(
        doc,
        [
            "Usar el boton de avance de cada tarjeta para mover el caso a la siguiente etapa.",
            "Si una accion requiere un dato obligatorio, el sistema abre un modal antes de confirmar.",
            "Si se cancela una accion obligatoria, el caso queda en su etapa anterior.",
        ],
    )

    doc.add_heading("Bandeja Tabla y filtros", level=2)
    doc.add_paragraph(
        "La vista tabla sirve para busquedas puntuales, control masivo y filtros por estado, sector, service externo o ubicacion."
    )
    add_image(doc, "05_bandeja_tabla_filtros.png", "Bandeja en formato tabla con filtros disponibles.")

    doc.add_heading("Detalle del caso", level=2)
    doc.add_paragraph(
        "El detalle concentra informacion del ticket, comentarios, historial, datos del rack, ubicaciones y adjuntos."
    )
    add_image(doc, "06_detalle_caso.png", "Detalle del caso: datos principales, comentarios e historial.")
    add_image(doc, "07_detalle_ubicaciones.png", "Detalle del caso: grilla de ubicaciones.")
    add_image(doc, "08_detalle_adjuntos.png", "Detalle del caso: adjuntos y fotos.")
    add_image(doc, "09_modal_accion.png", "Modal de accion: confirma datos obligatorios antes de avanzar.")

    doc.add_heading("5. Trabajo por rol", level=1)
    doc.add_heading("ADO", level=2)
    add_bullets(
        doc,
        [
            "Ingresar a la Bandeja Kanban y revisar los casos en la columna de ADO.",
            "Abrir el detalle y controlar zona, pasillo, cara, ubicaciones, niveles y descripcion de rotura.",
            "Si las posiciones son validas, cargar el numero de service externo.",
            "Si falta informacion o la posicion no existe, registrar el reclamo o rechazo con observacion clara.",
        ],
    )
    doc.add_heading("Mapa", level=2)
    add_bullets(
        doc,
        [
            "Tomar los casos derivados por ADO con service externo informado.",
            "Registrar observaciones de traspasos WMS cuando corresponda.",
            "Confirmar ubicaciones vacias antes de inutilizar en WMS.",
            "Luego de mantenimiento, relevar fisicamente, reetiquetar si corresponde y rehabilitar en WMS.",
        ],
    )
    doc.add_heading("Mantenimiento", level=2)
    add_bullets(
        doc,
        [
            "Trabajar los casos recibidos con service externo ya generado.",
            "Ejecutar la reparacion del rack.",
            "Al finalizar, derivar a Mapa para relevo fisico y cierre operativo.",
        ],
    )
    doc.add_heading("Ingenieria / soporte", level=2)
    add_bullets(
        doc,
        [
            "Monitorear Dashboard y Control de Ubicaciones.",
            "Revisar Ingresos Forms con error tecnico o reclamo pendiente.",
            "Acompanhar diferencias entre VigIA y WMS hasta su resolucion.",
        ],
    )

    doc.add_heading("6. Estados, acciones y reglas", level=1)
    add_matrix(
        doc,
        ["Estado", "Accion esperada", "Regla principal"],
        [
            ("Registrado", "ADO revisa la solicitud.", "Debe validar datos de ubicacion antes de avanzar."),
            ("Pendiente Service Externo", "ADO carga el service externo.", "No puede pasar a Mapa sin service externo."),
            ("Pendiente Traspasos", "Mapa analiza y registra observaciones.", "No requiere cargar detalle puntual de traspaso."),
            ("Traspasos Finalizados", "Mapa confirma ubicaciones vacias.", "Si no estan vacias, el caso permanece en control."),
            ("Posicion Bloqueada", "Mantenimiento ejecuta la reparacion.", "La ubicacion debe estar inutilizada en WMS."),
            ("En Reparacion", "Mantenimiento finaliza y deriva.", "Debe quedar trazabilidad en historial."),
            ("Reparado", "Mapa releva fisicamente.", "Si requiere reetiquetar, se registra antes del cierre."),
            ("Cerrado", "Caso finalizado.", "Mapa debe quitar inutilizacion WMS y cerrar."),
        ],
        (1.55, 2.45, 2.5),
    )
    add_callout(
        doc,
        "Regla critica",
        "Un caso no debe avanzar desde ADO hacia Mapa si no tiene cargado el numero de service externo.",
        color=RED,
    )

    doc.add_heading("7. Ingresos desde Forms", level=1)
    doc.add_paragraph(
        "La solapa Ingresos Forms muestra los archivos JSON recibidos desde el flujo de Power Automate. "
        "Desde alli se puede revisar el estado de importacion, reintentar o reclamar cuando el ingreso no cumple con los datos minimos."
    )
    add_image(doc, "10_ingresos_forms.png", "Ingresos Forms: monitoreo de JSON importados.")
    add_bullets(
        doc,
        [
            "OK: el JSON genero un caso en VigIA.",
            "Error tecnico: revisar formato del archivo, permisos, ruta o adjuntos.",
            "Reclamo: el operador debe corregir o completar informacion de origen.",
            "Reintentar: vuelve a procesar el ingreso cuando el problema ya fue resuelto.",
        ],
    )

    doc.add_heading("8. Buenas practicas y soporte", level=1)
    add_bullets(
        doc,
        [
            "Usar comentarios operativos breves, claros y accionables.",
            "Evitar avanzar estados sin confirmar la condicion real de WMS o del rack.",
            "Revisar adjuntos/fotos desde la solapa correspondiente, no desde texto pegado en comentarios.",
            "Usar filtros por service externo o ubicacion para investigar tickets puntuales.",
            "Ante diferencias WMS vs VigIA, revisar primero Control de Ubicaciones y luego el historial del caso.",
        ],
    )
    add_metadata_table(
        doc,
        [
            ("Soporte funcional", "Ingenieria / usuarios administradores de VigIA"),
            ("Evidencia sugerida", "Numero de caso, service externo, ubicacion, captura y descripcion del problema"),
            ("Ruta de consulta", "Dashboard -> Bandeja -> Detalle -> Historial / Adjuntos"),
        ],
        widths=(1.875, 4.625),
    )

    polish_spanish(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
