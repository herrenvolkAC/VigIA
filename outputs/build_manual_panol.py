from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import shutil

base = Path(r'C:\Ingenieria\VigIA\outputs\template_panol_insumos.docx')
out = Path(r'C:\Ingenieria\VigIA\outputs\Manual_Usuario_Panol_Insumos.docx')
flow_img = Path(r'C:\Ingenieria\VigIA\outputs\Panol_Insumos_Flujos_Modelo_Logisu.png')
shutil.copyfile(base, out)
doc = Document(out)
body = doc._body._element
for child in list(body):
    if child.tag.endswith('}sectPr'):
        continue
    body.remove(child)
section = doc.sections[0]
section.top_margin = Inches(1.05)
section.bottom_margin = Inches(0.85)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.different_first_page_header_footer = False
sectPr = section._sectPr
pgNumType = sectPr.find(qn('w:pgNumType'))
if pgNumType is None:
    pgNumType = OxmlElement('w:pgNumType')
    sectPr.append(pgNumType)
pgNumType.set(qn('w:start'), '1')
try:
    ht = section.header.tables[0]
    ht.cell(0,2).text = 'CÓDIGO: CD-00.00.000'
    ht.cell(1,2).text = 'VIGENCIA: 29/06/2026'
    ht.cell(2,2).text = 'REVISIÓN: 00'
    ht.cell(2,1).text = 'Pañol de Insumos'
    ht.cell(3,1).text = 'Pañol de Insumos'
    page_cell = ht.cell(3,2)
    page_cell.text = ''
    pnum = page_cell.paragraphs[0]
    pnum.add_run('PÁGINA: ')
    def add_field(paragraph, instr):
        run = paragraph.add_run()
        fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
        instr_text = OxmlElement('w:instrText'); instr_text.set(qn('xml:space'), 'preserve'); instr_text.text = instr
        fld_sep = OxmlElement('w:fldChar'); fld_sep.set(qn('w:fldCharType'), 'separate')
        text = OxmlElement('w:t'); text.text = '1'
        fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_begin); run._r.append(instr_text); run._r.append(fld_sep); run._r.append(text); run._r.append(fld_end)
    add_field(pnum, ' PAGE ')
    pnum.add_run(' de ')
    add_field(pnum, ' NUMPAGES ')
except Exception:
    pass
styles = doc.styles
for style_name in ['Normal', 'List Paragraph']:
    if style_name in styles:
        st = styles[style_name]
        st.font.name = 'Arial'
        st.font.size = Pt(10)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.08
for style_name in ['Heading 1', 'Heading 2', 'Heading 4']:
    if style_name in styles:
        st = styles[style_name]
        st.font.name = 'Arial'
        st.font.color.rgb = RGBColor(31, 78, 121)
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(5)
if 'Heading 1' in styles:
    styles['Heading 1'].font.size = Pt(12)
    styles['Heading 1'].font.bold = True
if 'Heading 2' in styles:
    styles['Heading 2'].font.size = Pt(11)
    styles['Heading 2'].font.bold = True
if 'Heading 4' in styles:
    styles['Heading 4'].font.size = Pt(10)
    styles['Heading 4'].font.bold = True
LIGHT_BLUE = 'D9EAF7'; LIGHT_GRAY = 'F2F2F2'; GREEN = 'E2F0D9'; AMBER = 'FFF2CC'

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ''; p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text); run.font.name = 'Arial'; run.font.size = Pt(size); run.bold = bold
    if color: run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_table_borders(table, color='A6A6A6'):
    tblPr = table._tbl.tblPr; borders = tblPr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders'); tblPr.append(borders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        element = borders.find(qn('w:' + edge))
        if element is None:
            element = OxmlElement('w:' + edge); borders.append(element)
        element.set(qn('w:val'), 'single'); element.set(qn('w:sz'), '4'); element.set(qn('w:space'), '0'); element.set(qn('w:color'), color)

def set_cell_margins(table, top=80, start=100, bottom=80, end=100):
    tblPr = table._tbl.tblPr; tblCellMar = tblPr.first_child_found_in('w:tblCellMar')
    if tblCellMar is None:
        tblCellMar = OxmlElement('w:tblCellMar'); tblPr.append(tblCellMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tblCellMar.find(qn('w:' + m))
        if node is None:
            node = OxmlElement('w:' + m); tblCellMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')

def add_heading(text, level=1):
    if level == 4:
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.italic = True
        r.font.color.rgb = RGBColor(31, 78, 121)
        return p
    return doc.add_paragraph(text, style=f'Heading {level}')

def add_para(text, style='Normal', bold_label=None):
    p = doc.add_paragraph(style=style)
    if bold_label:
        r = p.add_run(bold_label); r.bold = True; r.font.name = 'Arial'; r.font.size = Pt(10); p.add_run(' ')
    r = p.add_run(text); r.font.name = 'Arial'; r.font.size = Pt(10)
    return p

def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        r = p.add_run('• ')
        r.font.name = 'Arial'
        r.font.size = Pt(10)
        r2 = p.add_run(item)
        r2.font.name = 'Arial'
        r2.font.size = Pt(10)

def add_numbered(items):
    for idx, item in enumerate(items, 1):
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        r = p.add_run(f'{idx}. ')
        r.font.name = 'Arial'
        r.font.size = Pt(10)
        r.bold = True
        r2 = p.add_run(item)
        r2.font.name = 'Arial'
        r2.font.size = Pt(10)

def add_table(headers, rows, widths=None, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers)); table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit = False
    set_table_borders(table); set_cell_margins(table)
    for i, h in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], header_fill); set_cell_text(table.rows[0].cells[i], h, bold=True, size=8.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row): set_cell_text(cells[i], str(value), size=8.5)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths): row.cells[i].width = Inches(width)
    doc.add_paragraph('')
    return table

def add_callout(title, body, fill=AMBER):
    table = doc.add_table(rows=1, cols=1); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color='B7B7B7'); set_cell_margins(table, top=100, bottom=100, start=140, end=140)
    cell = table.cell(0,0); set_cell_shading(cell, fill); cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(title + ': '); r.bold = True; r.font.name = 'Arial'; r.font.size = Pt(9)
    r2 = p.add_run(body); r2.font.name = 'Arial'; r2.font.size = Pt(9)
    doc.add_paragraph('')

add_heading('OBJETIVO', 1)
add_para('Establecer la forma de uso del módulo Pañol de Insumos de VigIA para administrar stock de insumos, registrar producción, gestionar pedidos de sectores y consultar indicadores operativos.')
add_heading('ALCANCE', 1)
add_para('El presente instructivo aplica a usuarios solicitantes, operadores y administradores que intervienen en el circuito de Pañol de Insumos.')
add_bullets(['Gestión de maestro de artículos/PLU y usos asociados.','Importación de stock CD e inventarios de Oficina ADO.','Registro de movimientos de stock de insumos.','Alta de producción, entrega de producción y consulta de stock producido.','Solicitud, confirmación y cierre de pedidos de insumos o producción.','Consulta de indicadores por stock, producción y pedidos.'])
add_callout('Premisa operativa', 'El stock de insumos y el stock producido son independientes. Al registrar producción no se descuentan recetas ni materias primas del stock de insumos.', fill=GREEN)
add_heading('DEFINICIONES Y ABREVIATURAS', 1)
add_table(['Término', 'Definición'], [['PLU','Código de artículo utilizado para identificar insumos o producción.'],['Stock de insumos','Stock operativo compuesto por CD, Jaula y Oficina ADO.'],['Stock producido','Saldo generado por altas de producción menos entregas de producción.'],['Día logístico','Rango operativo desde el día anterior a las 22:00 hasta la fecha y hora actual.'],['Turno','Valor calculado automáticamente por el sistema según la hora del movimiento.'],['Uso','Clasificación opcional del artículo que, si existe, debe informarse al entregar un pedido.'],['Pedido parcial','Pedido satisfecho parcialmente y cerrado. Si queda necesidad pendiente se genera un nuevo pedido.']], widths=[1.45,5.75])
add_heading('RESPONSABILIDADES', 1)
add_table(['Rol','Responsabilidades principales'], [['Solicitante','Cargar pedidos para su sector, consultar stock disponible y revisar pedidos anteriores.'],['Operador','Registrar movimientos, producción, entregas, inventarios y satisfacer pedidos pendientes.'],['Administrador','Mantener artículos, importar stock CD, depurar datos operativos y supervisar indicadores.'],['Ingeniería / Sistemas','Mantener la configuración del módulo y ejecutar ajustes de base de datos desde backend cuando corresponda.']], widths=[1.8,5.4])
add_heading('FRECUENCIA', 1)
add_table(['Actividad','Frecuencia sugerida'], [['Importación stock CD','Cada vez que se disponga de archivo actualizado REP_STK_CD_UNID o equivalente.'],['Inventario Oficina ADO','Por turno o según criterio operativo definido por el área.'],['Movimientos de insumos','Cada vez que se produzca alta, baja, ajuste o transferencia.'],['Producción y entrega','En el momento de producir o entregar unidades a un sector.'],['Pedidos de insumos','A demanda de cada sector solicitante.'],['Revisión de indicadores','Diaria o al inicio de cada jornada operativa.']], widths=[2.2,5.0])
add_heading('DESARROLLO', 1)
add_heading('Ingreso al módulo y perfiles', 2)
add_numbered(['Ingresar a VigIA con usuario habilitado.','Seleccionar el módulo Pañol de Insumos desde el selector de aplicaciones.','Verificar que el perfil visualizado corresponda al rol de trabajo: Solicitante, Operador o Administrador.'])
add_table(['Perfil','Solapas disponibles','Acciones destacadas'], [['Solicitante','Pedido de Insumos','Genera solicitudes y consulta pedidos propios.'],['Operador','Indicadores, Stock Insumos, Producción, Historial, Pedidos Pendientes','Opera stock, producción, inventarios y cierre de pedidos.'],['Administrador','Todas las solapas','Además mantiene artículos, importa CD y ejecuta depuración operativa.']], widths=[1.4,2.9,2.9])
add_heading('Indicadores', 2)
add_para('La solapa Indicadores es el tablero inicial del módulo. Presenta tres columnas: Stock, Producción y Pedidos.')
add_numbered(['Definir el rango Fecha desde / Fecha hasta. Por defecto se muestra desde el primer día del mes actual hasta la fecha actual.','Opcionalmente seleccionar un PLU para filtrar ambos conceptos: insumos, producción y pedidos.','Presionar Aplicar para actualizar el tablero.','Usar Exportar Excel para descargar movimientos de producción por PLU y sector según el filtro aplicado.'])
add_table(['Bloque','Qué muestra','Uso operativo'], [['Stock','Artículos, bajo mínimo, movimientos, stock por ubicación y consumo por turno.','Detectar faltantes y riesgos de cobertura.'],['Producción','Producido, entregado, stock producido, entregas por sector y producción por turno.','Evaluar producción y malgastos por sector/turno.'],['Pedidos','Pedidos, pendientes, confirmados, pedidos por sector y top PLUs solicitados.','Controlar demanda y carga de trabajo pendiente.']], widths=[1.3,3.1,2.8])
add_heading('Stock Insumos', 2)
add_para('La solapa Stock Insumos concentra la operación del inventario de insumos. El stock de insumos no se mezcla con el stock producido.')
add_heading('Stock actual', 4); add_bullets(['Permite consultar stock CD, Jaula, Oficina ADO, total, mínimo, estado y cobertura.','El estado Bajo mínimo se calcula comparando stock total contra stock mínimo del artículo.','Los datos se actualizan luego de importaciones, inventarios y movimientos.'])
add_heading('Inventario', 4); add_numbered(['Seleccionar fecha y ubicación. Para Oficina ADO se usa por defecto la ubicación operativa principal.','Ingresar stock físico por PLU activo.','Guardar inventario. El sistema registra usuario, fecha, turno y calcula consumos cuando corresponde.'])
add_heading('Movimientos', 4); add_numbered(['Seleccionar artículo/PLU y tipo de movimiento.','Informar origen y/o destino según corresponda al tipo de movimiento.','Ingresar cantidad, motivo y observación si aplica.','Registrar el movimiento. El sistema valida que no se genere stock negativo cuando descuenta de una ubicación.'])
add_table(['Tipo de movimiento','Requiere origen','Requiere destino','Impacto'], [['ALTA / AJUSTE_POSITIVO','No','Sí','Incrementa stock destino.'],['BAJA / AJUSTE_NEGATIVO','Sí','No','Disminuye stock origen.'],['TRANSFERENCIA','Sí','Sí','Resta del origen y suma al destino.']], widths=[2.1,1.2,1.2,2.7])
add_heading('Artículos', 2)
add_para('La solapa Artículos permite mantener el maestro de PLUs utilizado tanto para insumos como para producción.')
add_numbered(['Ingresar código, descripción, categoría, unidad y stock mínimo.','Completar Uso solo si el artículo requiere clasificar el destino de uso al entregar pedidos.','Guardar el artículo. Los cambios impactan en los combos y grillas luego de actualizar la pantalla.'])
add_callout('Uso asociado','Si un PLU tiene lista de usos, el operador deberá seleccionar un uso al confirmar la entrega de un pedido que incluya ese PLU.', fill=AMBER)
add_heading('Importación CD y depuración operativa', 2)
add_heading('Importar archivo de stock CD', 4); add_numbered(['Ingresar a Importación CD.','Seleccionar archivo de stock CD, por ejemplo REP_STK_CD_UNID.csv.','Usar Vista previa para validar coincidencias de códigos propios.','Ejecutar Importar CD para actualizar stock CD de artículos existentes.']); add_bullets(['Los PLUs no existentes en el maestro propio se ignoran.','La importación no crea artículos automáticamente.','El resultado informa códigos coincidentes e ignorados.'])
add_heading('Depuración operativa', 4); add_para('La depuración borra datos operativos para iniciar pruebas o puesta en marcha sin eliminar PLUs.'); add_bullets(['Requiere clave autorizada.','Borra stock CD, movimientos, inventarios, consumos, producción y pedidos.','Conserva maestro de artículos, ubicaciones y turnos.'])
add_heading('Producción', 2)
add_para('La solapa Producción permite registrar unidades producidas y entregas de stock producido a sectores.')
add_heading('Alta de producción', 4); add_numbered(['Seleccionar PLU producido.','Ingresar cantidad producida.','Completar observación si corresponde.','Registrar producción. El sistema guarda usuario y calcula turno automáticamente según la hora.'])
add_heading('Entrega a sector', 4); add_numbered(['Seleccionar PLU entregado. El combo muestra solo PLUs con stock producido disponible.','Seleccionar sector destino.','Ingresar cantidad entregada.','Registrar entrega. El sistema descuenta del stock producido.'])
add_table(['Control','Regla'], [['Sector productor','No se solicita por pantalla; solo se almacena usuario que realiza la producción.'],['Turno','Se calcula automáticamente por hora en producción y entrega.'],['Stock producido','Se calcula como producción acumulada menos entregas acumuladas.'],['Indicadores','Permiten analizar producción por turno y entregas por sector/PLU.']], widths=[1.8,5.4])
add_heading('Pedido de Insumos', 2)
add_para('La solapa Pedido de Insumos está orientada al perfil solicitante. Permite pedir cantidades de insumo y/o producción para un sector.')
add_numbered(['Seleccionar sector solicitante.','Revisar la grilla de PLUs con stock positivo. La grilla muestra stock de insumo y stock de producción en modo solo lectura.','Ingresar cantidad a pedir en columna Insumo y/o Producción.','Completar observación si corresponde.','Enviar solicitud. El pedido queda en estado Pendiente y se notifica a operadores/administradores.']); add_bullets(['No se permite pedir más del stock disponible de cada tipo.','El solicitante puede consultar pedidos anteriores desde la misma solapa.','El solicitante no confirma la entrega; esa acción corresponde a operador o administrador.'])
add_heading('Pedidos Pendientes', 2)
add_para('La solapa Pedidos Pendientes permite a operadores y administradores satisfacer solicitudes de los sectores.')
add_numbered(['Seleccionar un pedido de la lista de pendientes.','Verificar sector, usuario solicitante, observación y líneas solicitadas.','Confirmar cantidades de insumo y/o producción a entregar. Puede ser una entrega parcial.','Para insumos, seleccionar origen. Por defecto se propone Oficina ADO.','Si el PLU tiene usos asociados, seleccionar el uso de entrega.','Confirmar entrega o cancelar pedido según corresponda.'])
add_callout('Entrega parcial','Si se entrega menos de lo solicitado, el pedido queda cerrado como confirmado parcial. Para solicitar faltantes debe generarse un nuevo pedido.', fill=AMBER)
add_table(['Concepto','Impacto al confirmar'], [['Insumo confirmado','Genera movimiento de baja desde la ubicación origen seleccionada.'],['Producción confirmada','Genera entrega de producción al sector solicitante y descuenta stock producido.'],['Uso de entrega','Se almacena en la línea del pedido cuando el PLU requiere uso.'],['Estado del pedido','Cambia a Confirmado o Confirmado Parcial y deja de figurar como pendiente.']], widths=[1.9,5.3])
add_heading('Historial', 2)
add_para('La solapa Historial permite consultar inventarios registrados por fecha, ubicación, turno y artículo. Se utiliza para auditoría operativa y revisión de consumos calculados.')
add_heading('Controles operativos recomendados', 2)
add_table(['Control','Responsable sugerido','Evidencia en sistema'], [['Revisar bajo mínimo al inicio de jornada','Operador / Administrador','Indicadores > Stock.'],['Controlar pedidos pendientes','Operador','Pedidos Pendientes e Indicadores > Pedidos.'],['Analizar entregas por sector','Administrador / Jefatura','Indicadores > Producción y exportación Excel.'],['Depurar datos de prueba antes de inicio oficial','Administrador autorizado','Importación CD > Depurar operativo.']], widths=[2.6,2.0,2.6])
add_heading('REGISTROS', 1)
add_table(['Registro','Ubicación / consulta'], [['Movimientos de insumos','Solapa Stock Insumos > Movimientos.'],['Inventarios por turno','Solapa Historial.'],['Movimientos de producción','Solapa Producción e Indicadores.'],['Pedidos de insumos','Solapas Pedido de Insumos y Pedidos Pendientes.'],['Exportación de indicadores de producción','Indicadores > Exportar Excel.']], widths=[2.6,4.6])
add_heading('ANEXOS', 1)
add_para('Anexo I: Mapa macro de flujos del módulo Pañol de Insumos.')
if flow_img.exists():
    doc.add_page_break(); add_heading('ANEXO I - MAPA MACRO DE FLUJOS', 1); add_para('El siguiente esquema resume los tres flujos principales: Stock de Insumos, Producción y Pedido/Entrega de Insumos.')
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(flow_img), width=Inches(7.2))
add_heading('REFERENCIAS', 1)
add_bullets(['Sistema VigIA - Módulo Pañol de Insumos.','Mapa macro de flujos: Panol_Insumos_Flujos_Modelo_Logisu.vsdx.','Criterios operativos definidos para stock de insumos, producción y pedidos.'])
add_heading('HISTORIAL DE CAMBIOS', 1)
add_table(['Revisión','Descripción del cambio','Fecha'], [['00','Emisión inicial del manual de uso del módulo Pañol de Insumos.','29/06/2026'],['','',''],['','','']], widths=[1.0,4.8,1.4], header_fill=LIGHT_GRAY)
props = doc.core_properties; props.title='Manual de uso - Pañol de Insumos'; props.subject='VigIA - Pañol de Insumos'; props.keywords='Pañol, Insumos, VigIA, Manual de uso, Producción, Pedidos'; props.comments='Generado a partir de template SGI provisto por el usuario.'
doc.save(out)
print(out)



