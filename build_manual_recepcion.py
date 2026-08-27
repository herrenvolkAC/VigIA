from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"C:\Ingenieria\VigIA")
OUT = ROOT / "docs" / "Manual_Usuario_Recepcion_Refrigerados.docx"
IMG_SELECTOR = ROOT / "docs" / "screenshots" / "selector_actual.png"
IMG_INGRESO = ROOT / "docs" / "screenshots" / "recepcion_ingreso_actual.png"
IMG_REGISTROS = ROOT / "docs" / "screenshots" / "recepcion_registros_actual.png"
IMG_DETALLE = ROOT / "docs" / "screenshots" / "recepcion_detalle_actual.png"

BLUE, NAVY, GREEN = "2E74B5", "0E1620", "1F7A4D"
LIGHT_BLUE, LIGHT_GREEN, GRAY, RULE = "E8EEF5", "EAF4EE", "5C6773", "D4DBD8"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd")) or OxmlElement("w:shd")
    if shd.getparent() is None:
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar") or OxmlElement("w:tcMar")
    if margins.getparent() is None:
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}")) or OxmlElement(f"w:{side}")
        if node.getparent() is None:
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def cell_border(cell, color=RULE, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders") or OxmlElement("w:tcBorders")
    if borders.getparent() is None:
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}")) or OxmlElement(f"w:{edge}")
        if node.getparent() is None:
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:tblHeader"))


def page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))


def caption(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string(GRAY)


def image_alt(inline_shape, text):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", text)
    doc_pr.set("descr", text)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.25
    p.add_run(text)


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5); p.paragraph_format.line_spacing = 1.25
    p.add_run(text)


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5); sec.page_height = Inches(11)
sec.top_margin = Inches(.75); sec.bottom_margin = Inches(.7)
sec.left_margin = Inches(.85); sec.right_margin = Inches(.85)
sec.header_distance = Inches(.35); sec.footer_distance = Inches(.35)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.font.color.rgb = RGBColor.from_string(NAVY)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in (
    ("Title", 26, NAVY, 0, 8), ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7), ("Heading 3", 12, "1F4D78", 10, 5),
):
    style = doc.styles[name]
    style.font.name = "Calibri"; style.font.size = Pt(size); style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = name != "Title"; style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after)

header = sec.header.paragraphs[0]; header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = header.add_run("VigIA · OPERACIÓN"); hr.font.size = Pt(9); hr.font.bold = True; hr.font.color.rgb = RGBColor.from_string(GREEN)
footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run("Manual de usuario · Recepción de refrigerados · Página "); fr.font.size = Pt(9); fr.font.color.rgb = RGBColor.from_string(GRAY)
page_number(footer)

p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
r = p.add_run("GUÍA OPERATIVA"); r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(GREEN)
title = doc.add_paragraph(style="Title"); title.add_run("Recepción de refrigerados")
sub = doc.add_paragraph(); sub.paragraph_format.space_after = Pt(18)
sr = sub.add_run("Manual de usuario para ingreso, consulta y seguimiento de descargas"); sr.font.size = Pt(13); sr.font.color.rgb = RGBColor.from_string(GRAY)

callout = doc.add_table(rows=1, cols=1); callout.alignment = WD_TABLE_ALIGNMENT.LEFT; callout.autofit = False
c = callout.cell(0, 0); c.width = Inches(6.65); shade(c, LIGHT_GREEN); cell_border(c, "A9C9B3", "8"); cell_margins(c, 140, 180, 140, 180)
cp = c.paragraphs[0]; cp.paragraph_format.space_after = Pt(3); cr = cp.add_run("Objetivo"); cr.bold = True; cr.font.color.rgb = RGBColor.from_string(GREEN)
bp = c.add_paragraph("Registrar cada descarga de productos refrigerados, asociar sus PLU, porcentajes de afectación y fotografías, y consultar posteriormente el detalle almacenado."); bp.paragraph_format.space_after = Pt(0)

doc.add_heading("1. Acceso al módulo", level=1)
doc.add_paragraph("Ingresar al sitio interno de VigIA con el usuario habilitado. Desde el selector de aplicaciones, elegir el módulo correspondiente a Recepción.")
if IMG_SELECTOR.exists():
    image_alt(doc.add_picture(str(IMG_SELECTOR), width=Inches(6.65)), "Selector de aplicaciones de VigIA")
    caption(doc, "Figura 1. Selector de aplicaciones de VigIA.")

doc.add_heading("2. Nueva descarga", level=1)
doc.add_paragraph("La solapa Nueva descarga permite cargar toda la información correspondiente a un ingreso. Los campos marcados con asterisco son obligatorios.")
doc.add_heading("2.1 Datos de la descarga", level=2)
for text in (
    "Fecha y hora: confirmar el momento de la recepción.",
    "Proveedor: escribir código o nombre y seleccionar una opción del listado Oracle.",
    "Recepcionista: buscar por legajo o nombre y seleccionar una opción del listado.",
    "Pallets recibidos y pallets auditados: ingresar cantidades numéricas. Los auditados no pueden superar a los recibidos.",
    "Cuenta con novedad y Observación: completar cuando corresponda.",
): bullet(doc, text)

doc.add_heading("2.2 Selección de artículos / PLU", level=2)
doc.add_paragraph("En PLU asociados, comenzar a escribir el código o la descripción del artículo. El sistema consulta Oracle y muestra hasta 50 coincidencias de los sectores habilitados para refrigerados.")
for text in (
    "Seleccionar el artículo desde la lista sugerida; no escribir solamente un texto libre.",
    "Completar el porcentaje de afectación del PLU cuando corresponda.",
    "Agregar una observación específica del artículo si es necesario.",
    "Usar + Agregar PLU para incorporar más artículos a la misma descarga.",
): bullet(doc, text)

doc.add_heading("2.3 Fotografías", level=2)
doc.add_paragraph("El módulo admite fotografías generales de la descarga y fotografías específicas de cada PLU.")
for text in (
    "Fotos de inspección: seleccionar imágenes generales del ingreso.",
    "Fotos de este PLU: utilizar el selector ubicado dentro de la fila del artículo para asociar imágenes a ese PLU.",
    "Se pueden seleccionar múltiples archivos. Esperar a que finalice la carga antes de cerrar o cambiar de pantalla.",
): bullet(doc, text)
if IMG_INGRESO.exists():
    image_alt(doc.add_picture(str(IMG_INGRESO), width=Inches(6.65)), "Pantalla actual de ingreso de recepción")
    caption(doc, "Figura 2. Pantalla actual de ingreso de una descarga.")

doc.add_heading("2.4 Guardar la descarga", level=2)
for text in ("Revisar proveedor, recepcionista, cantidades, PLU, porcentajes y fotografías.", "Presionar Guardar descarga.", "Esperar el mensaje de confirmación con el ID asignado."): numbered(doc, text)

doc.add_page_break()
doc.add_heading("3. Consultar registros", level=1)
doc.add_paragraph("La solapa Registros permite buscar descargas ya guardadas. Se puede filtrar por proveedor o PLU y por rango de fechas.")
if IMG_REGISTROS.exists():
    image_alt(doc.add_picture(str(IMG_REGISTROS), width=Inches(6.65)), "Pantalla actual de registros de recepción")
    caption(doc, "Figura 3. Pantalla actual de registros.")
for text in ("Ingresar un criterio en el filtro o seleccionar fechas desde/hasta.", "Presionar Buscar.", "En la grilla se visualizan fecha, proveedor, recepcionista, pallets, cantidad de PLU y cantidad de fotos.", "Presionar Abrir para acceder al detalle."): numbered(doc, text)

doc.add_heading("4. Detalle de una descarga", level=1)
doc.add_paragraph("El detalle reúne los datos principales del ingreso, los PLU asociados y las fotografías.")
if IMG_DETALLE.exists():
    image_alt(doc.add_picture(str(IMG_DETALLE), width=Inches(6.65)), "Pantalla actual del detalle de una descarga")
    caption(doc, "Figura 4. Detalle actual con fotos y porcentajes por PLU.")
for text in ("La información de cada PLU incluye código, descripción, porcentaje de afectación y sus fotos asociadas.", "Las fotos generales aparecen separadas de las fotos específicas por PLU.", "Presionar Cerrar para volver a la lista de registros. En tablet, el botón permanece visible en el encabezado del detalle."): bullet(doc, text)

doc.add_heading("5. Recomendaciones de uso", level=1)
table = doc.add_table(rows=1, cols=2); table.alignment = WD_TABLE_ALIGNMENT.LEFT; table.autofit = False
table.columns[0].width = Inches(1.6); table.columns[1].width = Inches(5.05)
for i, text in enumerate(("Situación", "Recomendación")):
    cell = table.cell(0, i); shade(cell, LIGHT_BLUE); cell_border(cell); cell_margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    run = cell.paragraphs[0].add_run(text); run.bold = True; run.font.color.rgb = RGBColor.from_string(NAVY)
repeat_header(table.rows[0])
for left, right in (
    ("Búsquedas", "Escribir al menos dos caracteres y esperar la lista de Oracle antes de continuar."),
    ("Fotos", "Evitar seleccionar archivos innecesariamente grandes o repetir la misma imagen."),
    ("Tablet", "Si el contenido excede la pantalla, desplazarse dentro del detalle; el encabezado y Cerrar permanecen disponibles."),
    ("Validación", "Antes de guardar, comprobar que cada PLU fue seleccionado desde la lista y que el porcentaje es correcto."),
):
    cells = table.add_row().cells
    for i, text in enumerate((left, right)):
        cell_border(cells[i]); cell_margins(cells[i]); cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP; cells[i].paragraphs[0].add_run(text)

doc.add_heading("6. Datos y almacenamiento", level=1)
doc.add_paragraph("Los maestros de proveedores, artículos y legajos se consultan en Oracle. Los datos propios de las recepciones se guardan en la base SQLite independiente del módulo y las fotografías se almacenan en la carpeta configurada del backend.")
note = doc.add_table(rows=1, cols=1); note.alignment = WD_TABLE_ALIGNMENT.LEFT; note.autofit = False
nc = note.cell(0, 0); shade(nc, "FFF7E6"); cell_border(nc, "E3B341", "8"); cell_margins(nc, 130, 160, 130, 160)
np = nc.paragraphs[0]; nr = np.add_run("Importante: "); nr.bold = True; np.add_run("no eliminar, mover ni renombrar manualmente la base SQLite o la carpeta de fotos. Las modificaciones de estructura se realizan desde el backend.")

end = doc.add_paragraph(); end.alignment = WD_ALIGN_PARAGRAPH.RIGHT
er = end.add_run("Versión de referencia: agosto de 2026"); er.italic = True; er.font.size = Pt(9); er.font.color.rgb = RGBColor.from_string(GRAY)
doc.core_properties.title = "Manual de usuario - Recepción de refrigerados"
doc.core_properties.subject = "Guía operativa del módulo VigIA"
doc.core_properties.author = "VigIA"
doc.core_properties.keywords = "VigIA, recepción, refrigerados, PLU, fotografías"
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
